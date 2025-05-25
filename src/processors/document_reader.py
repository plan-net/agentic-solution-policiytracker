import os
import tempfile
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from io import BytesIO

import pypdf
import markdown
from bs4 import BeautifulSoup
from docx import Document
import structlog

from src.models.content import DocumentType, DocumentMetadata
from src.utils.exceptions import DocumentProcessingError
from src.utils.validators import (
    validate_file_size,
    validate_file_extension,
    get_supported_extensions,
)
from src.config import settings
from src.integrations.azure_storage import AzureStorageClient

logger = structlog.get_logger()


class DocumentReader:
    """Multi-format document reader with Azure Storage support."""

    def __init__(self):
        self.supported_extensions = get_supported_extensions()
        self.max_size_mb = settings.MAX_DOCUMENT_SIZE_MB

    async def read_document(
        self, file_path: str, use_azure: bool = False
    ) -> Tuple[str, DocumentMetadata]:
        """Read document and extract text content and metadata from local or Azure Storage."""
        try:
            if use_azure:
                return await self._read_document_azure(file_path)
            else:
                return await self._read_document_local(file_path)

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(file_path, f"Unexpected error: {str(e)}")

    async def _read_document_local(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Read document from local filesystem."""
        # Validate file
        if not os.path.exists(file_path):
            raise DocumentProcessingError(file_path, "File does not exist")

        if not validate_file_extension(file_path, self.supported_extensions):
            raise DocumentProcessingError(file_path, "Unsupported file format")

        if not validate_file_size(file_path, self.max_size_mb):
            raise DocumentProcessingError(file_path, f"File too large (>{self.max_size_mb}MB)")

        # Determine document type
        extension = os.path.splitext(file_path)[1].lower()
        doc_type = self._get_document_type(extension)

        # Extract metadata
        metadata = self._extract_base_metadata(file_path, doc_type)

        # Extract text based on type
        if doc_type == DocumentType.PDF:
            text, extraction_meta = await self._read_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            text, extraction_meta = await self._read_docx(file_path)
        elif doc_type in [DocumentType.MARKDOWN]:
            text, extraction_meta = await self._read_markdown(file_path)
        elif doc_type == DocumentType.HTML:
            text, extraction_meta = await self._read_html(file_path)
        elif doc_type == DocumentType.TXT:
            text, extraction_meta = await self._read_text(file_path)
        else:
            raise DocumentProcessingError(file_path, f"No reader for type: {doc_type}")

        # Update metadata with extraction info
        metadata.extraction_metadata = extraction_meta

        # Validate extracted text
        if not text or len(text.strip()) < 10:
            raise DocumentProcessingError(file_path, "No meaningful text content extracted")

        # Truncate if too long (1MB text limit)
        if len(text) > 1024 * 1024:
            logger.warning("Text content truncated", file_path=file_path, original_length=len(text))
            text = text[: 1024 * 1024] + "\n\n[Content truncated due to size limit]"
            metadata.extraction_metadata["truncated"] = True

        return text, metadata

    async def _read_document_azure(self, blob_name: str) -> Tuple[str, DocumentMetadata]:
        """Read document from Azure Storage."""
        azure_client = AzureStorageClient()

        # Download blob content
        blob_data = await azure_client.download_blob("input-documents", blob_name)
        if not blob_data:
            raise DocumentProcessingError(blob_name, "Blob not found in Azure Storage")

        # Check size
        if len(blob_data) > self.max_size_mb * 1024 * 1024:
            raise DocumentProcessingError(blob_name, f"Blob too large (>{self.max_size_mb}MB)")

        # Determine document type from blob name
        extension = Path(blob_name).suffix.lower()
        if not validate_file_extension(blob_name, self.supported_extensions):
            raise DocumentProcessingError(blob_name, "Unsupported file format")

        doc_type = self._get_document_type(extension)

        # Get blob properties for metadata
        blob_properties = await azure_client.get_blob_properties("input-documents", blob_name)
        metadata = self._extract_azure_metadata(
            blob_name, doc_type, blob_properties, len(blob_data)
        )

        # Extract text based on type using temporary file for complex formats
        if doc_type in [DocumentType.PDF, DocumentType.DOCX]:
            # Create temporary file for binary formats
            with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as tmp_file:
                tmp_file.write(blob_data)
                tmp_file_path = tmp_file.name

            try:
                if doc_type == DocumentType.PDF:
                    text, extraction_meta = await self._read_pdf(tmp_file_path)
                elif doc_type == DocumentType.DOCX:
                    text, extraction_meta = await self._read_docx(tmp_file_path)
            finally:
                os.unlink(tmp_file_path)  # Clean up temp file

        else:
            # For text-based formats, process directly from bytes
            text_content = blob_data.decode("utf-8")

            if doc_type == DocumentType.MARKDOWN:
                text, extraction_meta = await self._read_markdown_content(text_content)
            elif doc_type == DocumentType.HTML:
                text, extraction_meta = await self._read_html_content(text_content)
            elif doc_type == DocumentType.TXT:
                text, extraction_meta = await self._read_text_content(text_content)
            else:
                raise DocumentProcessingError(blob_name, f"No reader for type: {doc_type}")

        # Update metadata with extraction info
        metadata.extraction_metadata = extraction_meta

        # Validate extracted text
        if not text or len(text.strip()) < 10:
            raise DocumentProcessingError(blob_name, "No meaningful text content extracted")

        # Truncate if too long (1MB text limit)
        if len(text) > 1024 * 1024:
            logger.warning("Text content truncated", file_path=blob_name, original_length=len(text))
            text = text[: 1024 * 1024] + "\n\n[Content truncated due to size limit]"
            metadata.extraction_metadata["truncated"] = True

        return text, metadata

    def _get_document_type(self, extension: str) -> DocumentType:
        """Map file extension to document type."""
        mapping = {
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
        }
        return mapping.get(extension, DocumentType.TXT)

    def _extract_base_metadata(self, file_path: str, doc_type: DocumentType) -> DocumentMetadata:
        """Extract basic file metadata."""
        stat = os.stat(file_path)

        return DocumentMetadata(
            source=os.path.basename(file_path),
            type=doc_type,
            file_path=file_path,
            file_size_bytes=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            extraction_metadata={},
        )

    async def _read_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Read PDF file using pypdf."""
        try:
            extraction_meta = {"pages_extracted": 0, "pages_failed": 0}
            text_parts = []

            with open(file_path, "rb") as file:
                reader = pypdf.PdfReader(file)

                # Check if encrypted
                if reader.is_encrypted:
                    raise DocumentProcessingError(file_path, "PDF is encrypted")

                total_pages = len(reader.pages)
                extraction_meta["total_pages"] = total_pages

                # Limit to 500 pages
                max_pages = min(total_pages, 500)

                for page_num in range(max_pages):
                    try:
                        page = reader.pages[page_num]
                        page_text = page.extract_text()

                        if page_text and page_text.strip():
                            text_parts.append(page_text)
                            extraction_meta["pages_extracted"] += 1
                        else:
                            extraction_meta["pages_failed"] += 1

                    except Exception as e:
                        logger.warning(
                            f"Failed to extract page {page_num}", file_path=file_path, error=str(e)
                        )
                        extraction_meta["pages_failed"] += 1

                if total_pages > 500:
                    extraction_meta["pages_truncated"] = total_pages - 500

            text = "\n\n".join(text_parts)
            return text, extraction_meta

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(file_path, f"PDF reading failed: {str(e)}")

    async def _read_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Read DOCX file using python-docx."""
        try:
            extraction_meta = {"paragraphs": 0, "tables": 0}
            text_parts = []

            doc = Document(file_path)

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    extraction_meta["paragraphs"] += 1

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    text_parts.append("\n".join(table_text))
                    extraction_meta["tables"] += 1

            text = "\n\n".join(text_parts)
            return text, extraction_meta

        except Exception as e:
            raise DocumentProcessingError(file_path, f"DOCX reading failed: {str(e)}")

    async def _read_markdown(self, file_path: str) -> Tuple[str, Dict]:
        """Read Markdown file and convert to plain text."""
        try:
            extraction_meta = {"format": "markdown"}

            with open(file_path, "r", encoding="utf-8") as file:
                md_content = file.read()

            # Convert markdown to HTML, then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text()

            # Preserve some structure
            text = text.replace("\n\n\n", "\n\n")  # Normalize line breaks

            extraction_meta["original_length"] = len(md_content)
            extraction_meta["converted_length"] = len(text)

            return text, extraction_meta

        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ["latin-1", "cp1252"]:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        md_content = file.read()

                    html = markdown.markdown(md_content)
                    soup = BeautifulSoup(html, "html.parser")
                    text = soup.get_text()

                    extraction_meta = {"format": "markdown", "encoding": encoding}
                    return text, extraction_meta

                except UnicodeDecodeError:
                    continue

            raise DocumentProcessingError(file_path, "Could not decode markdown file")

        except Exception as e:
            raise DocumentProcessingError(file_path, f"Markdown reading failed: {str(e)}")

    async def _read_html(self, file_path: str) -> Tuple[str, Dict]:
        """Read HTML file and extract text."""
        try:
            extraction_meta = {"format": "html"}

            with open(file_path, "r", encoding="utf-8") as file:
                html_content = file.read()

            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            extraction_meta["original_length"] = len(html_content)
            extraction_meta["converted_length"] = len(text)

            return text, extraction_meta

        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ["latin-1", "cp1252"]:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        html_content = file.read()

                    soup = BeautifulSoup(html_content, "html.parser")
                    for script in soup(["script", "style"]):
                        script.decompose()

                    text = soup.get_text()
                    lines = (line.strip() for line in text.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    text = "\n".join(chunk for chunk in chunks if chunk)

                    extraction_meta = {"format": "html", "encoding": encoding}
                    return text, extraction_meta

                except UnicodeDecodeError:
                    continue

            raise DocumentProcessingError(file_path, "Could not decode HTML file")

        except Exception as e:
            raise DocumentProcessingError(file_path, f"HTML reading failed: {str(e)}")

    async def _read_text(self, file_path: str) -> Tuple[str, Dict]:
        """Read plain text file."""
        try:
            extraction_meta = {"format": "text"}

            # Try UTF-8 first
            try:
                with open(file_path, "r", encoding="utf-8") as file:
                    text = file.read()
                extraction_meta["encoding"] = "utf-8"

            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                    try:
                        with open(file_path, "r", encoding=encoding) as file:
                            text = file.read()
                        extraction_meta["encoding"] = encoding
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise DocumentProcessingError(file_path, "Could not decode text file")

            # Basic cleanup
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            text = text.replace("\x00", "")  # Remove null bytes

            extraction_meta["original_length"] = len(text)

            return text, extraction_meta

        except Exception as e:
            raise DocumentProcessingError(file_path, f"Text reading failed: {str(e)}")

    def _extract_azure_metadata(
        self,
        blob_name: str,
        doc_type: DocumentType,
        blob_properties: Optional[Dict],
        file_size: int,
    ) -> DocumentMetadata:
        """Extract metadata from Azure blob properties."""
        now = datetime.now()

        # Use blob properties if available, otherwise use current time
        created_at = now
        modified_at = now

        if blob_properties:
            if "last_modified" in blob_properties:
                modified_at = blob_properties["last_modified"]
                created_at = modified_at  # Azure blobs don't have separate creation time
            elif "creation_time" in blob_properties:
                created_at = blob_properties["creation_time"]
                modified_at = blob_properties.get("last_modified", created_at)

        return DocumentMetadata(
            source=Path(blob_name).name,
            type=doc_type,
            file_path=blob_name,
            file_size_bytes=file_size,
            created_at=created_at,
            modified_at=modified_at,
            extraction_metadata={"source": "azure_storage"},
        )

    async def _read_markdown_content(self, content: str) -> Tuple[str, Dict]:
        """Read Markdown content from string."""
        try:
            extraction_meta = {"format": "markdown"}

            # Convert markdown to HTML, then extract text
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text()

            # Preserve some structure
            text = text.replace("\n\n\n", "\n\n")  # Normalize line breaks

            extraction_meta["original_length"] = len(content)
            extraction_meta["converted_length"] = len(text)

            return text, extraction_meta

        except Exception as e:
            raise DocumentProcessingError("markdown_content", f"Markdown reading failed: {str(e)}")

    async def _read_html_content(self, content: str) -> Tuple[str, Dict]:
        """Read HTML content from string."""
        try:
            extraction_meta = {"format": "html"}

            soup = BeautifulSoup(content, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text
            text = soup.get_text()

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            extraction_meta["original_length"] = len(content)
            extraction_meta["converted_length"] = len(text)

            return text, extraction_meta

        except Exception as e:
            raise DocumentProcessingError("html_content", f"HTML reading failed: {str(e)}")

    async def _read_text_content(self, content: str) -> Tuple[str, Dict]:
        """Read plain text content from string."""
        try:
            extraction_meta = {"format": "text", "encoding": "utf-8"}

            # Basic cleanup
            text = content.replace("\r\n", "\n").replace("\r", "\n")
            text = text.replace("\x00", "")  # Remove null bytes

            extraction_meta["original_length"] = len(text)

            return text, extraction_meta

        except Exception as e:
            raise DocumentProcessingError("text_content", f"Text reading failed: {str(e)}")
